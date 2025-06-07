from flask import Flask, request, render_template, send_file
import os
from openai import OpenAI
from docx import Document

app = Flask(__name__)
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

def create_docx(title, content, filename):
    doc = Document()
    doc.add_heading(title, 0)
    for line in content.strip().split("\n"):
        doc.add_paragraph(line)
    filepath = os.path.join("/mnt/data", filename)
    doc.save(filepath)
    return filepath

@app.route("/", methods=["GET"])
def form():
    return render_template("form.html")

@app.route("/questions", methods=["POST"])
def questions():
    task = request.form.get("task")

    prompt = f"Generate 20 tailored RAMS pre-task planning questions based on the task: {task}. Use UK civil engineering terminology."

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a RAMS generator that only outputs plain tailored questions."},
                {"role": "user", "content": prompt}
            ]
        )
        text = response.choices[0].message.content
        questions = [line.strip("1234567890. ").strip() for line in text.strip().split("\n") if line.strip()][:20]
        return render_template("questions.html", task=task, questions=questions)
    except Exception as e:
        return f"<h2>❌ Error generating questions: {str(e)}</h2>"

@app.route("/generate", methods=["POST"])
def generate():
    user_inputs = [request.form.get(f"q{i}") for i in range(1, 21)]
    formatted_input = "\n".join([f"{i+1}. {user_inputs[i]}" for i in range(20)])

    prompt = f"""
Create a full RAMS package based on these 20 user responses:

{formatted_input}

---

Generate the RAMS in 3 separate parts:
🟦 Stage 1 – Risk Assessment (minimum 550 words)
🟦 Stage 2 – Sequence of Activities (600+ words)
🟦 Stage 3 – Method Statement (750+ words)

Stage 3 must include these headings:
1. Scope of Works
2. Personnel and Responsibilities
3. Hold Points
4. Operated Plant
5. Non-Operated Plant
6. Materials Required
7. PPE Required
8. Rescue Plan (if applicable)
9. Applicable C2V+ Site Standards
10. CESWI References
11. Quality Control
12. Roles and Responsibilities
13. Environmental Considerations (5 paragraphs)

Each stage must follow that format and be output as separate plain text blocks.
Use UK construction terminology only.
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a specialist RAMS writer..."},
                {"role": "user", "content": prompt}
            ]
        )

        output = response.choices[0].message.content

        # ✅ Validate output structure
        if "Stage 2" not in output or "Stage 3" not in output:
            return "<h2>❌ Error: The response from GPT is missing required sections (Stage 2 or Stage 3).</h2>"

        parts = (
            output.split("Stage 2")[0],
            output.split("Stage 2")[1].split("Stage 3")[0],
            output.split("Stage 3")[1]
        )

        ra_path = create_docx("Stage 1 - Risk Assessment", parts[0], "Risk_Assessment.docx")
        sa_path = create_docx("Stage 2 - Sequence of Activities", parts[1], "Sequence_of_Activities.docx")
        ms_path = create_docx("Stage 3 - Method Statement", parts[2], "Method_Statement.docx")

        return f"""
        <h2>✅ RAMS generated successfully.</h2>
        <ul>
            <li><a href='/download?file=Risk_Assessment.docx'>Download Risk Assessment</a></li>
            <li><a href='/download?file=Sequence_of_Activities.docx'>Download Sequence of Activities</a></li>
            <li><a href='/download?file=Method_Statement.docx'>Download Method Statement</a></li>
        </ul>
        """

    except Exception as e:
        return f"<h2>❌ Error generating RAMS: {str(e)}</h2>"

@app.route("/download")
def download():
    file = request.args.get("file")
    return send_file(os.path.join("/mnt/data", file), as_attachment=True)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
