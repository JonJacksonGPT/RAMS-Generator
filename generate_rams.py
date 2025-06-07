from docx import Document
import os

def generate_all_rams(answers):
    paths = []

    doc1 = Document()
    doc1.add_heading("Stage 1 – Risk Assessment", level=1)
    table = doc1.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Hazard'
    hdr_cells[1].text = 'Person(s) at Risk'
    hdr_cells[2].text = 'Undesired Event'
    hdr_cells[3].text = 'Control Measures'
    hdr_cells[4].text = 'Actioned By'
    for i in range(20):
        row = table.add_row().cells
        row[0].text = f"Hazard {i+1}"
        row[1].text = "Operatives"
        row[2].text = "Injury"
        row[3].text = "Task-specific control"
        row[4].text = "Supervisor"
    path1 = "/tmp/Stage_1_Risk_Assessment.docx"
    doc1.save(path1)
    paths.append(path1)

    doc2 = Document()
    doc2.add_heading("Stage 2 – Sequence of Activities", level=1)
    doc2.add_paragraph("Sequence of activities generated based on user input:")
    for i, ans in enumerate(answers):
        doc2.add_paragraph(f"Step {i+1}: {ans}")
    path2 = "/tmp/Stage_2_Sequence_of_Activities.docx"
    doc2.save(path2)
    paths.append(path2)

    doc3 = Document()
    doc3.add_heading("Stage 3 – Method Statement", level=1)
    sections = [
        "1. Scope of Works", "2. Personnel and Responsibilities", "3. Hold Points",
        "4. Operated Plant", "5. Non-Operated Plant", "6. Materials Required",
        "7. PPE Required", "8. Rescue Plan", "9. Applicable Site Standards",
        "10. CESWI References", "11. Quality Control",
        "12. Roles and Qualifications", "13. Environmental Considerations"
    ]
    for i, section in enumerate(sections):
        doc3.add_heading(section, level=2)
        doc3.add_paragraph(answers[i] if i < len(answers) else "N/A")
    path3 = "/tmp/Stage_3_Method_Statement.docx"
    doc3.save(path3)
    paths.append(path3)

    return paths
